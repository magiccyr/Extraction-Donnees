import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import pandas as pd
from docx import Document
import re
from pathlib import Path

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
    from reportlab.lib import colors
except ImportError:
    SimpleDocTemplate = None


class UniversalConverterApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Convertisseur Universel PDF ⇄ Word ⇄ Excel")
        self.root.geometry("950x700")
        
        self.check_dependencies()
        self.setup_ui()
    
    def check_dependencies(self):
        """Vérifie que les modules nécessaires sont installés."""
        missing = []
        if pdfplumber is None:
            missing.append("pdfplumber")
        if PyPDF2 is None:
            missing.append("PyPDF2")
        if SimpleDocTemplate is None:
            missing.append("reportlab")
        
        if missing:
            messagebox.showwarning(
                "Modules manquants",
                f"Certains modules sont manquants:\n{', '.join(missing)}\n\n"
                f"Installez avec: pip install {' '.join(missing)}"
            )
    
    def setup_ui(self):
        main_frame = ttk.Frame(self.root, padding="10")
        main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        self.root.columnconfigure(0, weight=1)
        self.root.rowconfigure(0, weight=1)
        main_frame.columnconfigure(0, weight=1)
        main_frame.rowconfigure(0, weight=1)
        
        # Création des onglets
        self.notebook = ttk.Notebook(main_frame)
        self.notebook.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        # Onglets
        self.setup_pdf_to_excel_tab()
        self.setup_excel_to_pdf_tab()
        self.setup_pdf_to_word_tab()
        self.setup_word_to_pdf_tab()
        
        # Barre de statut
        self.status_label = ttk.Label(main_frame, text="Prêt", relief=tk.SUNKEN)
        self.status_label.grid(row=1, column=0, sticky=(tk.W, tk.E), pady=5)
    
    # ==================== PDF → EXCEL ====================
    def setup_pdf_to_excel_tab(self):
        frame = ttk.Frame(self.notebook, padding="10")
        self.notebook.add(frame, text="📊 PDF → Excel")
        
        frame.columnconfigure(1, weight=1)
        frame.rowconfigure(4, weight=1)
        
        ttk.Label(frame, text="Extracteur PDF vers Excel", 
                 font=('Arial', 14, 'bold')).grid(row=0, column=0, columnspan=3, pady=10)
        
        ttk.Label(frame, text="Fichier PDF:").grid(row=1, column=0, sticky=tk.W, pady=5)
        self.pdf_excel_entry = ttk.Entry(frame, width=50)
        self.pdf_excel_entry.grid(row=1, column=1, sticky=(tk.W, tk.E), pady=5, padx=5)
        ttk.Button(frame, text="Parcourir", 
                  command=self.browse_file_pdf_excel).grid(row=1, column=2, pady=5)
        
        ttk.Button(frame, text="Extraire les données", 
                  command=self.extract_pdf_data).grid(row=2, column=0, columnspan=3, pady=15)
        
        ttk.Label(frame, text="Aperçu des données extraites:", 
                 font=('Arial', 10, 'bold')).grid(row=3, column=0, columnspan=3, sticky=tk.W, pady=5)
        
        tree_frame = ttk.Frame(frame)
        tree_frame.grid(row=4, column=0, columnspan=3, sticky=(tk.W, tk.E, tk.N, tk.S), pady=5)
        
        vsb = ttk.Scrollbar(tree_frame, orient="vertical")
        hsb = ttk.Scrollbar(tree_frame, orient="horizontal")
        self.pdf_excel_tree = ttk.Treeview(tree_frame, yscrollcommand=vsb.set, 
                                           xscrollcommand=hsb.set, height=12)
        vsb.config(command=self.pdf_excel_tree.yview)
        hsb.config(command=self.pdf_excel_tree.xview)
        
        self.pdf_excel_tree.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        vsb.grid(row=0, column=1, sticky=(tk.N, tk.S))
        hsb.grid(row=1, column=0, sticky=(tk.W, tk.E))
        tree_frame.columnconfigure(0, weight=1)
        tree_frame.rowconfigure(0, weight=1)
        
        ttk.Button(frame, text="💾 Exporter vers Excel", 
                  command=self.export_to_excel).grid(row=5, column=0, columnspan=3, pady=15)
        
        self.pdf_excel_df = None
    
    # ==================== EXCEL → PDF ====================
    def setup_excel_to_pdf_tab(self):
        frame = ttk.Frame(self.notebook, padding="10")
        self.notebook.add(frame, text="📄 Excel → PDF")
        
        frame.columnconfigure(1, weight=1)
        frame.rowconfigure(3, weight=1)
        
        ttk.Label(frame, text="Excel vers PDF", 
                 font=('Arial', 14, 'bold')).grid(row=0, column=0, columnspan=3, pady=10)
        
        ttk.Label(frame, text="Fichier Excel:").grid(row=1, column=0, sticky=tk.W, pady=5)
        self.excel_pdf_entry = ttk.Entry(frame, width=50)
        self.excel_pdf_entry.grid(row=1, column=1, sticky=(tk.W, tk.E), pady=5, padx=5)
        ttk.Button(frame, text="Parcourir", 
                  command=self.browse_file_excel_pdf).grid(row=1, column=2, pady=5)
        
        ttk.Button(frame, text="Charger l'aperçu", 
                  command=self.load_excel_preview).grid(row=2, column=0, columnspan=3, pady=15)
        
        tree_frame = ttk.Frame(frame)
        tree_frame.grid(row=3, column=0, columnspan=3, sticky=(tk.W, tk.E, tk.N, tk.S), pady=5)
        
        vsb = ttk.Scrollbar(tree_frame, orient="vertical")
        hsb = ttk.Scrollbar(tree_frame, orient="horizontal")
        self.excel_pdf_tree = ttk.Treeview(tree_frame, yscrollcommand=vsb.set, 
                                           xscrollcommand=hsb.set, height=15)
        vsb.config(command=self.excel_pdf_tree.yview)
        hsb.config(command=self.excel_pdf_tree.xview)
        
        self.excel_pdf_tree.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        vsb.grid(row=0, column=1, sticky=(tk.N, tk.S))
        hsb.grid(row=1, column=0, sticky=(tk.W, tk.E))
        tree_frame.columnconfigure(0, weight=1)
        tree_frame.rowconfigure(0, weight=1)
        
        ttk.Button(frame, text="📄 Convertir vers PDF", 
                  command=self.convert_excel_to_pdf).grid(row=4, column=0, columnspan=3, pady=15)
        
        self.excel_pdf_df = None
    
    # ==================== PDF → WORD ====================
    def setup_pdf_to_word_tab(self):
        frame = ttk.Frame(self.notebook, padding="10")
        self.notebook.add(frame, text="📝 PDF → Word")
        
        frame.columnconfigure(1, weight=1)
        frame.rowconfigure(4, weight=1)
        
        ttk.Label(frame, text="Convertisseur PDF vers Word", 
                 font=('Arial', 14, 'bold')).grid(row=0, column=0, columnspan=3, pady=10)
        
        ttk.Label(frame, text="Fichier PDF:").grid(row=1, column=0, sticky=tk.W, pady=5)
        self.pdf_word_entry = ttk.Entry(frame, width=50)
        self.pdf_word_entry.grid(row=1, column=1, sticky=(tk.W, tk.E), pady=5, padx=5)
        ttk.Button(frame, text="Parcourir", 
                  command=self.browse_file_pdf_word).grid(row=1, column=2, pady=5)
        
        options_frame = ttk.LabelFrame(frame, text="Options de conversion", padding="10")
        options_frame.grid(row=2, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=10)
        
        self.pdf_word_formatting = tk.BooleanVar(value=True)
        ttk.Checkbutton(options_frame, text="Préserver le formatage (paragraphes)", 
                       variable=self.pdf_word_formatting).pack(anchor=tk.W)
        
        self.pdf_word_page_breaks = tk.BooleanVar(value=True)
        ttk.Checkbutton(options_frame, text="Ajouter des sauts de page", 
                       variable=self.pdf_word_page_breaks).pack(anchor=tk.W)
        
        ttk.Label(frame, text="Aperçu du contenu:", 
                 font=('Arial', 10, 'bold')).grid(row=3, column=0, columnspan=3, sticky=tk.W, pady=5)
        
        text_frame = ttk.Frame(frame)
        text_frame.grid(row=4, column=0, columnspan=3, sticky=(tk.W, tk.E, tk.N, tk.S), pady=5)
        
        text_scroll = ttk.Scrollbar(text_frame, orient="vertical")
        self.pdf_word_text = tk.Text(text_frame, height=15, width=70, 
                                     yscrollcommand=text_scroll.set, wrap=tk.WORD)
        text_scroll.config(command=self.pdf_word_text.yview)
        
        self.pdf_word_text.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        text_scroll.grid(row=0, column=1, sticky=(tk.N, tk.S))
        text_frame.columnconfigure(0, weight=1)
        text_frame.rowconfigure(0, weight=1)
        
        ttk.Button(frame, text="📝 Convertir vers Word", 
                  command=self.convert_pdf_to_word).grid(row=5, column=0, columnspan=3, pady=15)
    
    # ==================== WORD → PDF ====================
    def setup_word_to_pdf_tab(self):
        frame = ttk.Frame(self.notebook, padding="10")
        self.notebook.add(frame, text="📄 Word → PDF")
        
        frame.columnconfigure(1, weight=1)
        frame.rowconfigure(3, weight=1)
        
        ttk.Label(frame, text="Convertisseur Word vers PDF", 
                 font=('Arial', 14, 'bold')).grid(row=0, column=0, columnspan=3, pady=10)
        
        ttk.Label(frame, text="Fichier Word:").grid(row=1, column=0, sticky=tk.W, pady=5)
        self.word_pdf_entry = ttk.Entry(frame, width=50)
        self.word_pdf_entry.grid(row=1, column=1, sticky=(tk.W, tk.E), pady=5, padx=5)
        ttk.Button(frame, text="Parcourir", 
                  command=self.browse_file_word_pdf).grid(row=1, column=2, pady=5)
        
        ttk.Label(frame, text="Aperçu du contenu:", 
                 font=('Arial', 10, 'bold')).grid(row=2, column=0, columnspan=3, sticky=tk.W, pady=5)
        
        text_frame = ttk.Frame(frame)
        text_frame.grid(row=3, column=0, columnspan=3, sticky=(tk.W, tk.E, tk.N, tk.S), pady=5)
        
        text_scroll = ttk.Scrollbar(text_frame, orient="vertical")
        self.word_pdf_text = tk.Text(text_frame, height=18, width=70, 
                                     yscrollcommand=text_scroll.set, wrap=tk.WORD)
        text_scroll.config(command=self.word_pdf_text.yview)
        
        self.word_pdf_text.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        text_scroll.grid(row=0, column=1, sticky=(tk.N, tk.S))
        text_frame.columnconfigure(0, weight=1)
        text_frame.rowconfigure(0, weight=1)
        
        ttk.Button(frame, text="📄 Convertir vers PDF", 
                  command=self.convert_word_to_pdf).grid(row=4, column=0, columnspan=3, pady=15)
    
    # ==================== FONCTIONS PDF → EXCEL ====================
    def browse_file_pdf_excel(self):
        filename = filedialog.askopenfilename(
            title="Sélectionner un fichier PDF",
            filetypes=[("Fichiers PDF", "*.pdf"), ("Tous", "*.*")]
        )
        if filename:
            self.pdf_excel_entry.delete(0, tk.END)
            self.pdf_excel_entry.insert(0, filename)
            self.status_label.config(text=f"Fichier: {Path(filename).name}")
    
    def extract_text_pdfplumber(self, pdf_path):
        """Retourne le texte brut du PDF."""
        texts = []
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                try:
                    texts.append(page.extract_text() or "")
                except:
                    texts.append("")
        return "\n\n".join(texts)

    def extract_tables_pdfplumber(self, pdf_path):
        """Retourne une liste de DataFrames extraits comme tableaux."""
        dfs = []
        with pdfplumber.open(pdf_path) as pdf:
            for pageno, page in enumerate(pdf.pages, start=1):
                try:
                    tables = page.extract_tables()
                except:
                    tables = []
                for t in tables:
                    if not t:
                        continue
                    header = t[0]
                    rows = t[1:]
                    cleaned = [[cell if cell is not None else "" for cell in row] for row in rows]
                    try:
                        df = pd.DataFrame(cleaned, columns=header)
                    except:
                        df = pd.DataFrame(cleaned)
                    df["_source_page"] = pageno
                    dfs.append(df)
        return dfs

    def parse_key_values(self, text):
        """Cherche des paires clé: valeur dans le texte."""
        kv = {}
        pattern = re.compile(r"^\s*([A-Za-z0-9 _\-\u00C0-\u017F]{2,60})\s*[:=\-]\s*(.+)$")
        for line in text.splitlines():
            m = pattern.match(line)
            if m:
                key = re.sub(r"\s+", " ", m.group(1).strip())
                val = m.group(2).strip()
                kv[key] = kv.get(key, "") + (" | " + val if key in kv else val)
        return kv

    def extract_pdf_data(self):
        pdf_path = self.pdf_excel_entry.get()
        if not pdf_path:
            messagebox.showwarning("Attention", "Sélectionnez un fichier PDF")
            return
        
        if pdfplumber is None:
            messagebox.showerror("Erreur", "pdfplumber n'est pas installé")
            return
        
        try:
            self.status_label.config(text="Extraction en cours...")
            self.root.update()
            
            text = self.extract_text_pdfplumber(pdf_path)
            tables = self.extract_tables_pdfplumber(pdf_path)
            
            # Si on a des tableaux
            if tables:
                try:
                    df = pd.concat(tables, ignore_index=True, sort=False)
                except ValueError:
                    df = pd.DataFrame()
                    for t in tables:
                        df = pd.concat([df, t], ignore_index=True, sort=False)
                df["raw_text"] = text
                cols = [c for c in df.columns if c != "raw_text"] + ["raw_text"]
                self.pdf_excel_df = df[cols].reset_index(drop=True)
            else:
                # Sinon chercher des paires clé:valeur
                kv = self.parse_key_values(text)
                if kv:
                    df = pd.DataFrame([kv])
                    df["raw_text"] = text
                    cols = [c for c in df.columns if c != "raw_text"] + ["raw_text"]
                    self.pdf_excel_df = df[cols].reset_index(drop=True)
                else:
                    # Fallback: texte brut
                    self.pdf_excel_df = pd.DataFrame([{"raw_text": text}])
            
            self.display_dataframe(self.pdf_excel_tree, self.pdf_excel_df)
            self.status_label.config(
                text=f"Extraction réussie: {len(self.pdf_excel_df)} lignes, {len(self.pdf_excel_df.columns)} colonnes"
            )
        except Exception as e:
            messagebox.showerror("Erreur", f"Extraction échouée: {str(e)}")
            self.status_label.config(text="Erreur lors de l'extraction")
    
    def export_to_excel(self):
        if self.pdf_excel_df is None or self.pdf_excel_df.empty:
            messagebox.showwarning("Attention", "Aucune donnée à exporter")
            return
        
        filename = filedialog.asksaveasfilename(
            defaultextension=".xlsx",
            filetypes=[("Excel", "*.xlsx"), ("Tous", "*.*")]
        )
        if filename:
            try:
                self.pdf_excel_df.to_excel(filename, index=False, engine='openpyxl')
                messagebox.showinfo("Succès", 
                    f"Exporté vers: {filename}\n\n"
                    f"Lignes: {len(self.pdf_excel_df)}\n"
                    f"Colonnes: {len(self.pdf_excel_df.columns)}")
                self.status_label.config(text=f"Exporté: {Path(filename).name}")
            except Exception as e:
                messagebox.showerror("Erreur", f"Export échoué: {str(e)}")
    
    # ==================== FONCTIONS EXCEL → PDF ====================
    def browse_file_excel_pdf(self):
        filename = filedialog.askopenfilename(
            title="Sélectionner un fichier Excel",
            filetypes=[("Excel", "*.xlsx *.xls"), ("Tous", "*.*")]
        )
        if filename:
            self.excel_pdf_entry.delete(0, tk.END)
            self.excel_pdf_entry.insert(0, filename)
            self.status_label.config(text=f"Fichier: {Path(filename).name}")
    
    def load_excel_preview(self):
        excel_path = self.excel_pdf_entry.get()
        if not excel_path:
            messagebox.showwarning("Attention", "Sélectionnez un fichier Excel")
            return
        
        try:
            self.excel_pdf_df = pd.read_excel(excel_path)
            self.display_dataframe(self.excel_pdf_tree, self.excel_pdf_df)
            self.status_label.config(text=f"Chargé: {len(self.excel_pdf_df)} lignes")
        except Exception as e:
            messagebox.showerror("Erreur", f"Lecture échouée: {str(e)}")
    
    def convert_excel_to_pdf(self):
        if self.excel_pdf_df is None or SimpleDocTemplate is None:
            messagebox.showwarning("Attention", "Chargez un fichier Excel et vérifiez que reportlab est installé")
            return
        
        filename = filedialog.asksaveasfilename(
            defaultextension=".pdf",
            filetypes=[("PDF", "*.pdf"), ("Tous", "*.*")]
        )
        if not filename:
            return
        
        try:
            self.status_label.config(text="Conversion en cours...")
            self.root.update()
            
            doc = SimpleDocTemplate(filename, pagesize=A4)
            elements = []
            styles = getSampleStyleSheet()
            
            # Titre
            title = Paragraph(f"<b>Export Excel: {Path(self.excel_pdf_entry.get()).name}</b>", 
                            styles['Title'])
            elements.append(title)
            elements.append(Spacer(1, 0.3*inch))
            
            # Données en tableau
            data = [self.excel_pdf_df.columns.tolist()] + self.excel_pdf_df.head(100).values.tolist()
            
            # Calculer largeur des colonnes
            num_cols = len(self.excel_pdf_df.columns)
            col_widths = [min(2*inch, 6.5*inch/num_cols)] * num_cols
            
            table = Table(data, colWidths=col_widths, repeatRows=1)
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 9),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 10),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
                ('FONTSIZE', (0, 1), (-1, -1), 8),
            ]))
            
            elements.append(table)
            
            if len(self.excel_pdf_df) > 100:
                elements.append(Spacer(1, 0.2*inch))
                elements.append(Paragraph(
                    f"<i>Affichage limité à 100 lignes sur {len(self.excel_pdf_df)}</i>", 
                    styles['Normal']
                ))
            
            doc.build(elements)
            messagebox.showinfo("Succès", f"PDF créé: {filename}")
            self.status_label.config(text=f"Créé: {Path(filename).name}")
        except Exception as e:
            messagebox.showerror("Erreur", f"Conversion échouée: {str(e)}")
    
    # ==================== FONCTIONS PDF → WORD ====================
    def browse_file_pdf_word(self):
        filename = filedialog.askopenfilename(
            title="Sélectionner un fichier PDF",
            filetypes=[("PDF", "*.pdf"), ("Tous", "*.*")]
        )
        if filename:
            self.pdf_word_entry.delete(0, tk.END)
            self.pdf_word_entry.insert(0, filename)
            self.status_label.config(text=f"Fichier: {Path(filename).name}")
            self.load_pdf_preview(filename, self.pdf_word_text)
    
    def load_pdf_preview(self, pdf_path, text_widget):
        if PyPDF2 is None:
            text_widget.delete(1.0, tk.END)
            text_widget.insert(tk.END, "PyPDF2 n'est pas installé")
            return
        
        try:
            text_widget.delete(1.0, tk.END)
            with open(pdf_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                for i in range(min(2, len(reader.pages))):
                    text = reader.pages[i].extract_text()
                    text_widget.insert(tk.END, f"--- Page {i+1} ---\n{text}\n\n")
                if len(reader.pages) > 2:
                    text_widget.insert(tk.END, f"... et {len(reader.pages) - 2} page(s) supplémentaire(s)")
        except Exception as e:
            text_widget.insert(tk.END, f"Erreur: {str(e)}")
    
    def convert_pdf_to_word(self):
        pdf_path = self.pdf_word_entry.get()
        if not pdf_path:
            messagebox.showwarning("Attention", "Sélectionnez un fichier PDF")
            return
        
        if PyPDF2 is None:
            messagebox.showerror("Erreur", "PyPDF2 n'est pas installé")
            return
        
        filename = filedialog.asksaveasfilename(
            defaultextension=".docx",
            initialfile=Path(pdf_path).stem + ".docx",
            filetypes=[("Word", "*.docx"), ("Tous", "*.*")]
        )
        if not filename:
            return
        
        try:
            self.status_label.config(text="Conversion en cours...")
            self.root.update()
            
            with open(pdf_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                doc = Document()
                doc.add_heading(f'Conversion de: {Path(pdf_path).name}', 0)
                doc.add_paragraph()
                
                for i in range(len(reader.pages)):
                    if self.pdf_word_page_breaks.get() and i > 0:
                        doc.add_page_break()
                    
                    doc.add_heading(f'Page {i+1}', level=2)
                    text = reader.pages[i].extract_text()
                    
                    if self.pdf_word_formatting.get():
                        for para in text.split('\n'):
                            if para.strip():
                                doc.add_paragraph(para.strip())
                    else:
                        doc.add_paragraph(text)
                
                doc.save(filename)
            
            messagebox.showinfo("Succès", 
                f"PDF converti avec succès!\n\n"
                f"Document Word créé: {filename}\n"
                f"Pages converties: {len(reader.pages)}")
            self.status_label.config(text=f"Créé: {Path(filename).name}")
        except Exception as e:
            messagebox.showerror("Erreur", f"Conversion échouée: {str(e)}")
    
    # ==================== FONCTIONS WORD → PDF ====================
    def browse_file_word_pdf(self):
        filename = filedialog.askopenfilename(
            title="Sélectionner un fichier Word",
            filetypes=[("Word", "*.docx"), ("Tous", "*.*")]
        )
        if filename:
            self.word_pdf_entry.delete(0, tk.END)
            self.word_pdf_entry.insert(0, filename)
            self.status_label.config(text=f"Fichier: {Path(filename).name}")
            self.load_word_preview(filename, self.word_pdf_text)
    
    def load_word_preview(self, word_path, text_widget):
        try:
            text_widget.delete(1.0, tk.END)
            doc = Document(word_path)
            for i, para in enumerate(doc.paragraphs[:50]):
                if para.text.strip():
                    text_widget.insert(tk.END, para.text + "\n")
            if len(doc.paragraphs) > 50:
                text_widget.insert(tk.END, f"\n... et {len(doc.paragraphs) - 50} paragraphes supplémentaires")
        except Exception as e:
            text_widget.insert(tk.END, f"Erreur: {str(e)}")
    
    def convert_word_to_pdf(self):
        word_path = self.word_pdf_entry.get()
        if not word_path:
            messagebox.showwarning("Attention", "Sélectionnez un fichier Word")
            return
        
        if SimpleDocTemplate is None:
            messagebox.showerror("Erreur", "reportlab n'est pas installé")
            return
        
        filename = filedialog.asksaveasfilename(
            defaultextension=".pdf",
            initialfile=Path(word_path).stem + ".pdf",
            filetypes=[("PDF", "*.pdf"), ("Tous", "*.*")]
        )
        if not filename:
            return
        
        try:
            self.status_label.config(text="Conversion en cours...")
            self.root.update()
            
            doc_word = Document(word_path)
            doc_pdf = SimpleDocTemplate(filename, pagesize=A4)
            elements = []
            styles = getSampleStyleSheet()
            
            # Titre
            title = Paragraph(f"<b>{Path(word_path).name}</b>", styles['Title'])
            elements.append(title)
            elements.append(Spacer(1, 0.3*inch))
            
            # Paragraphes
            for para in doc_word.paragraphs:
                if para.text.strip():
                    # Détection de titres
                    if para.style.name.startswith('Heading'):
                        p = Paragraph(f"<b>{para.text}</b>", styles['Heading2'])
                    else:
                        # Échapper les caractères spéciaux pour reportlab
                        text = para.text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                        p = Paragraph(text, styles['Normal'])
                    elements.append(p)
                    elements.append(Spacer(1, 0.1*inch))
            
            doc_pdf.build(elements)
            messagebox.showinfo("Succès", f"PDF créé: {filename}")
            self.status_label.config(text=f"Créé: {Path(filename).name}")
        except Exception as e:
            messagebox.showerror("Erreur", f"Conversion échouée: {str(e)}")
    
    # ==================== AFFICHAGE ====================
    def display_dataframe(self, tree, df):
        """Affiche un DataFrame dans un Treeview."""
        # Effacer les données précédentes
        for item in tree.get_children():
            tree.delete(item)
        
        if df is None or df.empty:
            return
        
        # Configuration des colonnes
        columns = list(df.columns)
        tree['columns'] = columns
        tree['show'] = 'headings'
        
        # En-têtes
        for col in columns:
            tree.heading(col, text=col)
            # Largeur selon le type de colonne
            if col == "raw_text":
                tree.column(col, width=200)
            else:
                tree.column(col, width=150)
        
        # Données (limiter à 100 lignes pour l'affichage)
        for _, row in df.head(100).iterrows():
            values = []
            for col in columns:
                val = str(row[col])
                # Tronquer les valeurs trop longues
                if len(val) > 100:
                    val = val[:97] + "..."
                values.append(val)
            tree.insert('', tk.END, values=values)
        
        # Indicateur si plus de 100 lignes
        if len(df) > 100:
            tree.insert('', tk.END, values=["..."] * len(columns))


def main():
    root = tk.Tk()
    app = UniversalConverterApp(root)
    root.mainloop()


if __name__ == "__main__":
    main()
